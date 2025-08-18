
from flask import Flask, request, jsonify
import os
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

app = Flask(__name__)
DOC_DIR = './documents'

def ensure_dir():
    os.makedirs(DOC_DIR, exist_ok=True)

def md_to_docx(content, out_path):
    doc = Document()
    for line in content.splitlines():
        if line.strip() == '':
            doc.add_paragraph('')
        else:
            doc.add_paragraph(line)
    doc.save(out_path)

def md_to_pdf(content, out_path):
    c = canvas.Canvas(out_path, pagesize=letter)
    width, height = letter
    y = height - 40
    for line in content.splitlines():
        c.drawString(40, y, line)
        y -= 14
        if y < 40:
            c.showPage()
            y = height - 40
    c.save()

@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint for service monitoring"""
    return jsonify({'status': 'ok', 'service': 'document-generation'})

@app.route('/api/documents/generate', methods=['POST'])
def generate():
    """Generate documents in multiple formats from content"""
    try:
        data = request.json or {}
        base = data.get('filename_base', 'document')
        content = data.get('content', '')
        formats = data.get('formats', ['md'])
        
        print(f"📄 Document generation request: {base} in formats {formats}")
        
        ensure_dir()
        paths = {}
        
        if 'md' in formats:
            md_path = os.path.join(DOC_DIR, f"{base}.md")
            with open(md_path, 'w', encoding='utf-8') as f:
                f.write(content)
            paths['md'] = md_path
            print(f"✅ Created MD: {md_path}")
            
        if 'docx' in formats:
            docx_path = os.path.join(DOC_DIR, f"{base}.docx")
            md_to_docx(content, docx_path)
            paths['docx'] = docx_path
            print(f"✅ Created DOCX: {docx_path}")
            
        if 'pdf' in formats:
            pdf_path = os.path.join(DOC_DIR, f"{base}.pdf")
            md_to_pdf(content, pdf_path)
            paths['pdf'] = pdf_path
            print(f"✅ Created PDF: {pdf_path}")
            
        print(f"🎉 Document generation completed successfully")
        return jsonify({'success': True, 'paths': paths})
        
    except Exception as e:
        print(f"❌ Document generation error: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500

if __name__ == '__main__':
    print("🚀 Starting Document Generation Service...")
    print("📂 Document output directory: ./documents")
    print("🌐 Service available at: http://0.0.0.0:5001")
    print("🏥 Health check: http://0.0.0.0:5001/health")
    print("📄 Generate endpoint: POST http://0.0.0.0:5001/api/documents/generate")
    app.run(host='0.0.0.0', port=5001, debug=True)
